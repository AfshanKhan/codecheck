"""Word (.docx) reporter -- for pasting into an email or sharing with someone
who won't open a .md/.json file. Same content as the markdown reporter, styled
as headings/tables instead of Markdown syntax.
"""

from __future__ import annotations

from pathlib import Path

import docx
from docx.shared import Pt, RGBColor

from codecheck.models import Finding, ReviewReport, Severity

_SEVERITY_COLOR = {
    Severity.CRITICAL: RGBColor(0xB4, 0x00, 0x00),
    Severity.HIGH: RGBColor(0xC0, 0x3B, 0x00),
    Severity.MEDIUM: RGBColor(0xB4, 0x64, 0x00),
    Severity.LOW: RGBColor(0x1F, 0x5C, 0xA8),
    Severity.INFO: RGBColor(0x60, 0x60, 0x60),
}


def _add_finding_row(table, f: Finding) -> None:
    row_cells = table.add_row().cells
    row_cells[0].paragraphs[0].add_run(str(f.line_start))

    run_sev = row_cells[1].paragraphs[0].add_run(f.severity.value.upper())
    run_sev.bold = True
    run_sev.font.color.rgb = _SEVERITY_COLOR[f.severity]

    row_cells[2].paragraphs[0].add_run(f"{f.check_id} ({f.source})")
    row_cells[3].paragraphs[0].add_run(f.title)


def render_docx(report: ReviewReport) -> docx.Document:
    doc = docx.Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10.5)

    title = doc.add_heading("codecheck Review Report", level=0)
    title.alignment = 1

    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run(f"Generated: {report.generated_at.isoformat()}")
    run.italic = True

    doc.add_heading("Summary", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Shading Accent 1"
    summary_rows = [
        ("Repo", report.repo_path),
        ("Mode", report.mode),
        ("Base ref", report.base_ref or "-"),
        ("Head ref", report.head_ref or "-"),
        ("Tiers run", ", ".join(report.tiers_run) or "-"),
        ("Files reviewed", str(len(report.files_reviewed))),
        ("Duration", f"{report.duration_seconds:.1f}s"),
    ]
    counts = report.counts_by_severity()
    for s in (Severity.CRITICAL, Severity.HIGH, Severity.MEDIUM, Severity.LOW, Severity.INFO):
        if counts[s]:
            summary_rows.append((f"{s.value.title()} findings", str(counts[s])))
    for label, value in summary_rows:
        row_cells = table.add_row().cells
        row_cells[0].paragraphs[0].add_run(label).bold = True
        row_cells[1].paragraphs[0].add_run(value)

    doc.add_heading("Findings", level=1)
    if not report.findings:
        doc.add_paragraph("No findings.")
    else:
        for file_path, findings in sorted(report.by_file().items()):
            doc.add_heading(file_path, level=2)
            file_table = doc.add_table(rows=1, cols=4)
            file_table.style = "Table Grid"
            hdr = file_table.rows[0].cells
            hdr[0].paragraphs[0].add_run("Line").bold = True
            hdr[1].paragraphs[0].add_run("Severity").bold = True
            hdr[2].paragraphs[0].add_run("Check").bold = True
            hdr[3].paragraphs[0].add_run("Title").bold = True
            for f in sorted(findings, key=lambda x: (-x.severity.rank, x.line_start)):
                _add_finding_row(file_table, f)

    if report.skipped:
        doc.add_heading("Skipped", level=1)
        for entry in report.skipped:
            doc.add_paragraph(entry, style="List Bullet")

    return doc


def write_docx_report(report: ReviewReport, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    render_docx(report).save(str(output_path))
